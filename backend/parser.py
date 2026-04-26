from pathlib import Path
from pypdf import PdfReader
from docx import Document


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _pdf(path)
    if suffix == ".docx":
        return _docx(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".doc":
        raise ValueError(f"Legacy .doc not supported — convert {path.name} to .docx or .pdf")
    raise ValueError(f"Unsupported file type: {suffix}")


def _pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts).strip()


def _docx(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()
